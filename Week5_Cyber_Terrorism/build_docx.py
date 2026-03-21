# Builds report_week5.docx from scratch using python-docx
# Run once: python3 build_docx.py
# Output:   report_week5.docx

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)


def set_font(run, name="Calibri", size=11, bold=False, italic=False, color=None):
    run.font.name   = name
    run.font.size   = Pt(size)
    run.font.bold   = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_para(text="", size=11, bold=False, italic=False, align=None,
             space_before=0, space_after=6, keep_with_next=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before    = Pt(space_before)
    p.paragraph_format.space_after     = Pt(space_after)
    p.paragraph_format.line_spacing    = Pt(size * 1.15)
    p.paragraph_format.keep_with_next  = keep_with_next
    if align:
        p.alignment = align
    if text:
        run = p.add_run(text)
        set_font(run, size=size, bold=bold, italic=italic)
    return p


def add_heading(text, level=1):
    sizes   = {1: 13, 2: 11}
    space_b = {1: 10, 2: 6}
    p = add_para(text, size=sizes[level], bold=True,
                 space_before=space_b[level], space_after=3,
                 keep_with_next=True)
    if level == 1:
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"),   "single")
        bottom.set(qn("w:sz"),    "4")
        bottom.set(qn("w:space"), "2")
        bottom.set(qn("w:color"), "000000")
        pBdr.append(bottom)
        pPr.append(pBdr)
    return p


def add_mixed(parts, space_before=0, space_after=6):
    """parts = list of (text, bold, italic, mono)"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    p.paragraph_format.line_spacing = Pt(11 * 1.15)
    for text, bold, italic, mono in parts:
        run = p.add_run(text)
        run.font.name   = "Courier New" if mono else "Calibri"
        run.font.size   = Pt(10 if mono else 11)
        run.font.bold   = bold
        run.font.italic = italic
    return p


def add_table(headers, rows, col_widths=None, mono_cols=None):
    mono_cols = mono_cols or []
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style     = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    hrow = table.rows[0]
    for i, h in enumerate(headers):
        cell = hrow.cells[i]
        cell.paragraphs[0].clear()
        run = cell.paragraphs[0].add_run(h)
        set_font(run, size=10, bold=True)

    for ri, row in enumerate(rows):
        trow = table.rows[ri + 1]
        for ci, val in enumerate(row):
            cell = trow.cells[ci]
            cell.paragraphs[0].clear()
            run = cell.paragraphs[0].add_run(str(val))
            set_font(run, name="Courier New" if ci in mono_cols else "Calibri",
                     size=9 if ci in mono_cols else 10)

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    doc.add_paragraph()
    return table


def bullet(parts):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after  = Pt(3)
    p.paragraph_format.line_spacing = Pt(11 * 1.15)
    p.paragraph_format.left_indent  = Cm(0.5)
    for text, bold, italic, mono in parts:
        run = p.add_run(text)
        run.font.name   = "Courier New" if mono else "Calibri"
        run.font.size   = Pt(10 if mono else 11)
        run.font.bold   = bold
        run.font.italic = italic
    return p


def ref_item(num, parts):
    p = doc.add_paragraph()
    p.paragraph_format.space_after       = Pt(2)
    p.paragraph_format.left_indent       = Cm(0.8)
    p.paragraph_format.first_line_indent = Cm(-0.8)
    p.paragraph_format.line_spacing      = Pt(11 * 1.15)
    run = p.add_run(f"[{num}] ")
    set_font(run, size=10, bold=True)
    for text, bold, italic, url in parts:
        run = p.add_run(text)
        run.font.size   = Pt(10)
        run.font.bold   = bold
        run.font.italic = italic
        if url:
            run.font.color.rgb = RGBColor(0x00, 0x56, 0xD2)
    return p


# ── TITLE ────────────────────────────────────────────────────────────────────

add_para("Report Module 5", size=16, bold=True,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
add_para("Cyber Terrorism and Information Warfare", size=13,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
add_para("CS 475: Introduction to Computer Security", size=11,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
add_para("Mária Nyolcas  |  Prof: Asen Nikolov Grozdanov", size=11,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
add_para("March 2026", size=11,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)

p = doc.add_paragraph()
pPr = p._p.get_or_add_pPr()
pBdr = OxmlElement("w:pBdr")
bot = OxmlElement("w:bottom")
bot.set(qn("w:val"), "single"); bot.set(qn("w:sz"), "6")
bot.set(qn("w:space"), "1");    bot.set(qn("w:color"), "000000")
pBdr.append(bot); pPr.append(pBdr)
p.paragraph_format.space_after = Pt(6)

add_mixed([
    ("Abstract.  ", True, False, False),
    ("This report examines cyber terrorism and information warfare as distinct but "
     "overlapping threat categories, combining theoretical analysis with three practical "
     "exercises: threat actor mapping using MITRE ATT&CK, an anatomy of a documented "
     "disinformation campaign (Operation Secondary Infektion), and the design of a "
     "fictional attack scenario against a hospital network with paired countermeasures. "
     "The report connects these topics to vulnerability and scanning work from Week 4 "
     "and ends with a forward link to industrial espionage.",
     False, False, False),
], space_after=10)


# ── SECTION 1 ────────────────────────────────────────────────────────────────

add_heading("1  Introduction")

add_para(
    "Cyber terrorism and information warfare are two of the most debated concepts in "
    "contemporary security. The confusion is partly terminological: the word "
    "\u201cterrorism\u201d carries legal and political weight that shapes how incidents are "
    "classified, investigated, and attributed. A working definition that cuts through "
    "the noise comes from RAND Corporation: cyber terrorism consists of "
    "\u201cpremeditated, politically motivated attacks against information, computer "
    "systems, computer programs, and data which result in violence against "
    "non-combatant targets by sub-national groups or clandestine agents.\u201d The key "
    "ingredients are intent (political), target (civilian or non-combatant "
    "infrastructure), and effect (fear, coercion, or physical disruption through "
    "digital means).",
    space_after=6)

add_para(
    "This distinguishes cyber terrorism from two closely related categories. "
    "Cybercrime is financially motivated: ransomware gangs, payment card fraud, and "
    "cryptocurrency theft are canonical examples. Hacktivism is ideologically motivated "
    "but typically non-violent: website defacement and DDoS protests are disruptive but "
    "rarely cause lasting physical harm. The boundary between the three is blurry in "
    "practice \u2014 a ransomware attack on a hospital may simultaneously constitute "
    "cybercrime (extortion), potential terrorism (endangering patients), and, if "
    "state-sponsored, an act of warfare.",
    space_after=6)

add_para(
    "Information Warfare (IW) predates the internet by centuries; propaganda leaflets "
    "dropped from aircraft in WWI are textbook examples. What digital networks changed "
    "is scale, speed, and deniability. Modern IW doctrine covers six sub-domains: "
    "psychological operations (PSYOP), electronic warfare, cyber operations, deception, "
    "OSINT exploitation, and censorship or narrative control. This module sits at the "
    "intersection of all six, making it the most cross-cutting topic in the course so far.",
    space_after=6)

add_para(
    "For a computer science student, the relevance is concrete: the same scanning and "
    "exploitation techniques studied in Week 4 (Nmap, OpenVAS, CVE analysis) are the "
    "operational foundation for the cyber operations sub-domain of IW. The distinction "
    "is motivation and scale, not method.",
    space_after=10)


# ── SECTION 2 ────────────────────────────────────────────────────────────────

add_heading("2  Theoretical Framework")

add_heading("2.1  Actor Taxonomy", level=2)
add_para(
    "Understanding who conducts cyber terrorism and IW is as important as understanding "
    "the technical methods, because the actor type determines capability, persistence, "
    "and the likelihood of attribution. Four categories matter for threat modelling:",
    space_after=4)

bullet([("Nation-state APTs", True, False, False),
        (" operate with virtually unlimited resources, long planning horizons, and "
         "geopolitical objectives. MITRE ATT&CK documents dozens of these groups with "
         "their full TTP profiles.", False, False, False)])
bullet([("Non-state terrorist organisations", True, False, False),
        (" increasingly use cyber capabilities for recruitment, propaganda distribution, "
         "and cryptocurrency fundraising \u2014 less often for direct disruption.",
         False, False, False)])
bullet([("Hacktivist collectives", True, False, False),
        (" are loosely organised and ideologically driven; their campaigns often blur "
         "into IW territory when they amplify disinformation alongside DDoS attacks.",
         False, False, False)])
bullet([("Insider threats", True, False, False),
        (" are relevant because, as Week 2\u2019s access-control analysis showed, "
         "overly permissive role assignments create the conditions that insiders "
         "(and APTs using stolen credentials) exploit.", False, False, False)])

doc.add_paragraph()

add_heading("2.2  Landmark Case Studies", level=2)

add_mixed([
    ("Estonia 2007.  ", True, False, False),
    ("A wave of DDoS attacks paralysed Estonian government, banking, and media websites "
     "following a political dispute with Russia over the relocation of a Soviet-era war "
     "memorial. This is the first large-scale cyber attack against a nation-state\u2019s "
     "digital infrastructure and directly accelerated the creation of NATO CCDCOE in "
     "Tallinn. The key takeaway is that even technically unsophisticated attacks "
     "(volumetric DDoS) can cause strategic-level disruption if the target lacks "
     "redundancy and has not planned for the scenario.", False, False, False),
], space_after=6)

add_mixed([
    ("NotPetya (2017).  ", True, False, False),
    ("Initially disguised as ransomware, NotPetya was a wiper \u2014 malware designed "
     "to permanently destroy data with no recovery mechanism. Attributed to the Sandworm "
     "APT group (Russian GRU Unit 74455), it caused an estimated $10 billion in global "
     "damages by spreading through a compromised Ukrainian accounting software update "
     "(supply-chain initial access), then destroying backups and MBRs across hundreds "
     "of organisations worldwide. The takeaway: a cyber weapon can cause catastrophic "
     "collateral damage well beyond its intended geographic or organisational target. "
     "The CVSS 9.8 vulnerabilities studied in Week 4 (network-reachable, "
     "no-preconditions exploits) map directly to how NotPetya spread laterally once "
     "inside a network perimeter.", False, False, False),
], space_after=10)


# ── SECTION 3 ────────────────────────────────────────────────────────────────

add_heading("3  Lab Findings")

add_heading("3.1  Exercise 1: MITRE ATT&CK Threat Actor Mapping", level=2)

add_mixed([
    ("Using the ATT&CK Navigator, I built heat-map layers for Sandworm Team (APT44, "
     "Russian GRU Unit 74455) and Lazarus Group (APT38, North Korea RGB). Structured "
     "data is encoded in ", False, False, False),
    ("threat_actor_mapping.py", False, False, True),
    (".", False, False, False),
], space_after=4)

add_para("Five Sandworm techniques selected:", bold=True, space_after=3)

add_table(
    headers=["ID", "Tactic", "Technique / observation"],
    rows=[
        ["T1486",     "Impact (TA0040)",          "Data Encrypted for Impact \u2014 NotPetya/BadRabbit: encryption is cover for destruction, no recovery key exists"],
        ["T1059.003", "Execution (TA0002)",        "Windows Command Shell \u2014 cmd.exe stages wipers and disables AV using built-in tools"],
        ["T1078",     "Defence Evasion (TA0005)",  "Valid Accounts \u2014 stolen credentials for lateral movement, blends with normal traffic in SIEM"],
        ["T1071.001", "C2 (TA0011)",               "Web Protocols \u2014 HTTPS C2 blends with legitimate browsing, hard to block selectively"],
        ["T1561.002", "Impact (TA0040)",           "Disk Structure Wipe \u2014 Industroyer2 overwrites MBR; full reinstall required for recovery"],
    ],
    col_widths=[2.4, 3.6, 10.3],
    mono_cols=[0],
)

add_mixed([
    ("Comparison with Lazarus Group (3 shared techniques).  ", True, False, False),
    ("Both groups use T1486 (wiper/ransomware), T1078 (valid accounts), and T1071 "
     "(application layer C2). This overlap is instructive rather than diagnostic: these "
     "are among the most effective and widely documented techniques in ATT&CK. Shared "
     "TTPs do not imply shared origin. Attribution requires corroborating evidence: "
     "infrastructure reuse, malware code similarity, timing patterns relative to "
     "geopolitical events, and HUMINT. The overlap illustrates why confident public "
     "attribution is rare.", False, False, False),
], space_after=8)

add_heading("3.2  Exercise 2: Disinformation Analysis", level=2)

add_para(
    "I analysed Operation Secondary Infektion using the Stanford Internet Observatory / "
    "EU DisinfoLab joint report as my primary source. The operation ran from at least "
    "2014 to 2020, planted forged documents across 300+ platforms in 7 languages, and "
    "targeted audiences in Ukraine, Germany, France, Spain, and the UK.",
    space_after=4)

add_mixed([
    ("Campaign anatomy.  ", True, False, False),
    ("Fabricated content included forged EU and NATO official documents, fake news "
     "articles mimicking established outlets, and manipulated images. Laundering "
     "occurred through newly registered typosquatting domains. Amplification relied "
     "on coordinated cross-posting and circular citation between fake sites to create "
     "an illusion of independent corroboration.", False, False, False),
], space_after=4)

add_mixed([
    ("OSINT tool applied: WHOIS/DNS lookup (lookup.icann.org).  ", True, False, False),
    ("Key findings: (i) registration date preceded first article by two weeks only, "
     "consistent with purpose-built infrastructure; (ii) name server shared with two "
     "other fake outlets identified in the same investigation \u2014 the single most "
     "useful finding, linking three \u201cindependent\u201d sites to the same actor. "
     "Limitation: WHOIS alone cannot prove intent; privacy-shielded registration is "
     "legal and common.", False, False, False),
], space_after=4)

add_para("Red-flag indicator table (6 indicators):", bold=True, space_after=3)

add_table(
    headers=["#", "Indicator", "How to check", "Confidence"],
    rows=[
        ["1", "Coordinated posting timing",      "Botometer; cross-platform timestamp comparison",               "High"],
        ["2", "Domain registration anomaly",     "WHOIS lookup; compare reg. date vs. first article",           "High"],
        ["3", "Document metadata mismatch",      "exiftool / PDF properties; compare software locale",          "High"],
        ["4", "Cross-platform self-citation loop","archive.org; trace content back to single origin",           "Medium"],
        ["5", "Linguistic inconsistency",         "Back-translation; native speaker review",                    "Medium"],
        ["6", "Infrastructure clustering",        "Reverse-IP lookup; certificate transparency logs",           "High"],
    ],
    col_widths=[0.6, 4.2, 6.0, 2.5],
)

add_heading("3.3  Exercise 3: Attack Scenario and Countermeasures", level=2)

add_mixed([
    ("Target: ", True, False, False),
    ("Hospital / emergency medical services network.  ", False, False, False),
    ("Actor: ", True, False, False),
    ("Iron Caduceus (fictional nation-state APT, advanced, state-level resources, "
     "politically motivated coercion). Full structured output in ",
     False, False, False),
    ("attack_scenario.py", False, False, True),
    (".", False, False, False),
], space_after=4)

add_table(
    headers=["Step", "Stage", "Attacker action", "Countermeasure"],
    rows=[
        ["1",  "Reconnaissance",       "Passive OSINT: LinkedIn, job posts, conference slides",         "OSINT audit of own public footprint; social media policy"],
        ["2*", "Weaponisation",        "Spear-phish impersonating EHR vendor; macro dropper",           "Disable macros by GPO; email sandboxing"],
        ["3",  "Delivery",             "Email to IT admins and procurement; timed Monday AM",           "Phishing simulation; DMARC/DKIM/SPF"],
        ["4",  "Exploitation",         "Macro spawns PowerShell; in-memory LOTL execution",             "EDR monitors process lineage; PS script-block logging"],
        ["5",  "Installation",         "Scheduled task + registry Run key; internal recon",             "SIEM alerts on task creation; network segmentation"],
        ["6",  "C&C",                  "HTTPS C2, domain fronting, randomised beacon interval",         "TLS inspection; DNS threat intelligence feeds"],
        ["7",  "Actions on Objectives","Lateral movement to EHR; wiper destroys records and backups",  "Offline immutable backups (3-2-1); PAW for EHR admin"],
    ],
    col_widths=[0.8, 2.8, 5.8, 6.9],
)

add_mixed([
    ("* Weakest link: Step 2 (Weaponisation/Delivery).  ", True, False, False),
    ("A staff member who recognises the phishing email and reports it before opening "
     "the attachment breaks the entire chain before any code executes. At every "
     "subsequent step the attacker already has a foothold, making defence progressively "
     "more expensive. Investment in training and email filtering at this stage has the "
     "highest defensive return on investment.", False, False, False),
], space_after=10)


# ── SECTION 4 ────────────────────────────────────────────────────────────────

add_heading("4  Reflection and Critical Analysis")

add_para(
    "Studying cybercrime \u2014 as we did implicitly in Week 4 through CVE analysis \u2014 "
    "focuses on technical vulnerabilities: a specific version of vsftpd has a known "
    "backdoor, and the fix is to update it. The exploitation is impersonal and largely "
    "automated. Cyber terrorism and information warfare are different in a way that "
    "matters for how a computer scientist thinks about defence.",
    space_after=6)

add_para(
    "The technical underpinning is often identical. NotPetya\u2019s lateral spread used "
    "EternalBlue, the same SMB exploit class as the CVE-2007-2447 Samba vulnerability "
    "examined in Week 4; the CVSS vector (network-reachable, no credentials required) "
    "is the same category of risk. What changes is the adversary model: a ransomware "
    "gang will stop if the victim pays or if the exploit is patched. An APT acting on "
    "political orders does not stop. It adapts, waits, and comes back through a "
    "different vector. This fundamentally changes the defensive posture required: "
    "patching is still necessary, but it is not sufficient when the adversary has "
    "the patience and resources to find the next unpatched system.",
    space_after=6)

add_para(
    "The disinformation exercise reinforced a different angle: the most effective "
    "information operations are not technically sophisticated. Operation Secondary "
    "Infektion did not require any CVEs. It required patience, language skills, and "
    "an understanding of how human trust in media works. The OSINT indicators I "
    "documented (metadata, registration timing, shared infrastructure) are technically "
    "checkable by any analyst \u2014 the defence against information warfare is partly "
    "technical and partly a matter of critical reading skills that have nothing to do "
    "with network security.",
    space_after=6)

add_para(
    "The weakest link in the hospital scenario (spear-phishing) connects directly to "
    "the Principle of Least Privilege discussion in Week 2: the reason spear-phishing "
    "is so effective against high-privilege accounts is precisely because those accounts "
    "have been granted rights that far exceed daily task requirements. Tightening access "
    "reduces the blast radius when the phish succeeds.",
    space_after=6)

add_para(
    "Looking ahead, information warfare intersects with industrial espionage in a "
    "specific way: state-sponsored actors conducting IW campaigns often simultaneously "
    "collect intellectual property and trade secrets from the same targets, using the "
    "same infrastructure \u2014 the two objectives are pursued in a single operation.",
    space_after=10)


# ── REFERENCES ────────────────────────────────────────────────────────────────

add_heading("References")

# APA 7th: alphabetical by first author/org, hanging indent, no numbers
# ref_item reused for consistent formatting; numbers shown as visual anchors only
refs = [
    ([("Center for Internet Security. (2024). ", False, False, False),
      ("CIS Controls v8", False, True, False),
      (". https://www.cisecurity.org/controls", False, False, True)]),
    ([("Cybersecurity and Infrastructure Security Agency [CISA]. (2022). ", False, False, False),
      ("Alert AA22-057A: Destructive malware targeting organizations in Ukraine", False, True, False),
      (". https://www.cisa.gov/news-events/cybersecurity-advisories/aa22-057a", False, False, True)]),
    ([("European Union Agency for Cybersecurity [ENISA]. (2023). ", False, False, False),
      ("ENISA threat landscape 2023", False, True, False),
      (". https://www.enisa.europa.eu/topics/cyberthreats/enisa-threat-landscape", False, False, True)]),
    ([("Greenberg, A. (2019). ", False, False, False),
      ("Sandworm: A new era of cyberwar and the hunt for the Kremlin\u2019s most dangerous hackers", False, True, False),
      (". Doubleday.", False, False, False)]),
    ([("Lockheed Martin. (2022). ", False, False, False),
      ("Cyber kill chain", False, True, False),
      (". https://www.lockheedmartin.com/en-us/capabilities/cyber/cyber-kill-chain.html", False, False, True)]),
    ([("MITRE Corporation. (2024). ", False, False, False),
      ("ATT&CK enterprise matrix v15", False, True, False),
      (". https://attack.mitre.org", False, False, True)]),
    ([("National Institute of Standards and Technology [NIST]. (2018). ", False, False, False),
      ("Framework for improving critical infrastructure cybersecurity", False, True, False),
      (" (Version 1.1). https://www.nist.gov/cyberframework", False, False, True)]),
    ([("NATO Cooperative Cyber Defence Centre of Excellence [CCDCOE]. (2023). ", False, False, False),
      ("Tallinn manual 2.0 on the international law applicable to cyber operations", False, True, False),
      (". https://ccdcoe.org/research/tallinn-manual/", False, False, True)]),
    ([("Rid, T. (2020). ", False, False, False),
      ("Active measures: The secret history of disinformation and political warfare", False, True, False),
      (". Farrar, Straus and Giroux.", False, False, False)]),
    ([("Stanford Internet Observatory & EU DisinfoLab. (2020). ", False, False, False),
      ("Secondary infektion: A large-scale cross-platform influence operation", False, True, False),
      (". https://secondaryinfektion.org", False, False, True)]),
]

for i, parts in enumerate(refs, 1):
    ref_item(i, parts)


out = "/home/maria/Computer_Security/Week5_Cyber_Terrorism/report_week5.docx"
doc.save(out)
print(f"saved -> {out}")
