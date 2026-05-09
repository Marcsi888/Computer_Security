# Builds report_week12_13.docx from scratch using python-docx
# Run once: python build_docx.py
# Output:   report_week12_13.docx

import os
from docx import Document

_HERE = os.path.dirname(os.path.abspath(__file__))
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH


def add_para(doc, text="", size=11, bold=False, italic=False, align=None, after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = Pt(size * 1.15)
    if align:
        p.alignment = align
    if text:
        r = p.add_run(text)
        r.font.name = "Calibri"
        r.font.size = Pt(size)
        r.bold = bold
        r.italic = italic
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(text, style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = Pt(11 * 1.15)


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"

    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h

    for r_i, row in enumerate(rows, start=1):
        for c_i, value in enumerate(row):
            table.rows[r_i].cells[c_i].text = str(value)

    doc.add_paragraph()


def main() -> None:
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    add_para(doc, "Report Module 12 and 13", size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
    add_para(doc, "Emerging Topics: Blockchain, DApps, and E-commerce Security", size=13, align=WD_ALIGN_PARAGRAPH.CENTER, after=4)
    add_para(doc, "CS 475: Introduction to Computer Security", size=11, align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
    add_para(doc, "Maria Nyolcas", size=11, align=WD_ALIGN_PARAGRAPH.CENTER, after=8)

    add_para(doc, "Introduction", size=12, bold=True, after=4)
    add_para(
        doc,
        "This installment covers Part IV of the continuous project, focused on blockchain and decentralized applications (Lesson 1) and e-commerce security issues (Lesson 2).",
    )

    add_para(doc, "Part A - Theory and Guided Research", size=12, bold=True, after=4)
    add_bullet(doc, "Explain what a blockchain is and what trust problem it solves.")
    add_bullet(doc, "Differentiate blockchain infrastructure from DApp architecture.")
    add_bullet(doc, "List smart-contract vulnerabilities that are unique or amplified.")

    add_para(doc, "Part B - Hands-on Practical Work", size=12, bold=True, after=4)
    add_bullet(doc, "Implemented a toy blockchain in Python with proof-of-work and chain validation.")
    add_bullet(doc, "Tampered a mined block and observed integrity failure.")
    add_bullet(doc, "Attempted attacker repair by re-mining one block and observed downstream mismatch.")

    add_table(
        doc,
        ["Risk", "Impact", "Control"],
        [
            ("Card-not-present fraud", "Chargebacks and direct loss", "3-D Secure 2 and velocity checks"),
            ("Credential stuffing", "Account takeover", "MFA and adaptive rate limiting"),
            ("Supply-chain JS compromise", "Payment skimming", "CSP and dependency governance"),
        ],
    )

    add_para(doc, "Part C - Critical Reflection", size=12, bold=True, after=4)
    add_para(
        doc,
        "This installment ties directly to authentication, applied cryptography, malware, network attacks, and security models. The core lesson is that decentralized protocols reduce trust in operators but do not eliminate endpoint and software risk.",
    )

    add_para(doc, "Conclusion", size=12, bold=True, after=4)
    add_para(
        doc,
        "Blockchain improves tamper-evident integrity by construction, but secure deployment still depends on key management, secure coding, and operational controls.",
    )

    doc.save(os.path.join(_HERE, "report_week12_13.docx"))
    print("report_week12_13.docx written successfully.")


if __name__ == "__main__":
    main()
