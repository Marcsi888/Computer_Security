# Builds report_week3.docx from scratch using python-docx
# Run once: python3 build_docx.py
# Output:   report_week3.docx

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)


def set_font(run, name="Calibri", size=11, bold=False, italic=False, color=None):
    run.font.name  = name
    run.font.size  = Pt(size)
    run.font.bold  = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_para(text="", style="Normal", space_before=0, space_after=6, size=11,
             bold=False, italic=False, align=None, keep_with_next=False):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    p.paragraph_format.line_spacing = Pt(size * 1.15)
    if align:
        p.alignment = align
    if keep_with_next:
        p.paragraph_format.keep_with_next = True
    if text:
        run = p.add_run(text)
        set_font(run, size=size, bold=bold, italic=italic)
    return p


def add_heading(text, level=1):
    sizes   = {1: 13, 2: 11}
    space_b = {1: 10, 2: 6}
    p = add_para(text, size=sizes[level], bold=True,
                 space_before=space_b[level], space_after=3,
                 keep_with_next=True)
    if level == 1:
        # underline via bottom border
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"),   "single")
        bottom.set(qn("w:sz"),    "4")
        bottom.set(qn("w:space"), "2")
        bottom.set(qn("w:color"), "000000")
        pBdr.append(bottom)
        pPr.append(pBdr)
    return p


def add_mixed(parts, space_before=0, space_after=6):
    """
    parts = list of (text, bold, italic, mono)
    mono uses Courier New
    """
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    p.paragraph_format.line_spacing = Pt(11 * 1.15)
    for text, bold, italic, mono in parts:
        run = p.add_run(text)
        run.font.name   = "Courier New" if mono else "Calibri"
        run.font.size   = Pt(10 if mono else 11)
        run.font.bold   = bold
        run.font.italic = italic
    return p


def add_table(headers, rows, col_widths=None, mono_cols=None):
    mono_cols = mono_cols or []
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # header row
    hrow = table.rows[0]
    for i, h in enumerate(headers):
        cell = hrow.cells[i]
        cell.paragraphs[0].clear()
        run = cell.paragraphs[0].add_run(h)
        set_font(run, size=10, bold=True)

    # data rows
    for ri, row in enumerate(rows):
        trow = table.rows[ri + 1]
        for ci, val in enumerate(row):
            cell = trow.cells[ci]
            cell.paragraphs[0].clear()
            run = cell.paragraphs[0].add_run(str(val))
            set_font(run, name="Courier New" if ci in mono_cols else "Calibri",
                     size=9 if ci in mono_cols else 10)

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    doc.add_paragraph()   # spacing after table
    return table


def bullet(text_parts, level=0):
    """text_parts = list of (text, bold, italic, mono)"""
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after  = Pt(3)
    p.paragraph_format.line_spacing = Pt(11 * 1.15)
    p.paragraph_format.left_indent  = Cm(0.5 + level * 0.5)
    for text, bold, italic, mono in text_parts:
        run = p.add_run(text)
        run.font.name   = "Courier New" if mono else "Calibri"
        run.font.size   = Pt(10 if mono else 11)
        run.font.bold   = bold
        run.font.italic = italic
    return p


def ref_item(num, parts):
    p = doc.add_paragraph()
    p.paragraph_format.space_after       = Pt(2)
    p.paragraph_format.left_indent       = Cm(0.8)
    p.paragraph_format.first_line_indent = Cm(-0.8)
    p.paragraph_format.line_spacing      = Pt(11 * 1.15)
    run = p.add_run(f"[{num}] ")
    set_font(run, size=10, bold=True)
    for text, bold, italic, url in parts:
        run = p.add_run(text)
        run.font.size   = Pt(10)
        run.font.bold   = bold
        run.font.italic = italic
        if url:
            run.font.color.rgb = RGBColor(0x00, 0x56, 0xD2)
    return p



add_para("Report Module 3", size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER,
         space_before=0, space_after=2)
add_para("Network Scanning, Attack Classification, and Vulnerability Analysis",
         size=13, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=0, space_after=4)
add_para("CS 475: Introduction to Computer Security",
         size=11, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=0, space_after=2)
add_para("Mária Nyolcas  |  Prof: Asen Nikolov Grozdanov",
         size=11, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=0, space_after=2)
add_para("March 2026", size=11, align=WD_ALIGN_PARAGRAPH.CENTER,
         space_before=0, space_after=6)

# thin rule
p = doc.add_paragraph()
pPr = p._p.get_or_add_pPr()
pBdr = OxmlElement("w:pBdr")
bot = OxmlElement("w:bottom")
bot.set(qn("w:val"), "single"); bot.set(qn("w:sz"), "6")
bot.set(qn("w:space"), "1");    bot.set(qn("w:color"), "000000")
pBdr.append(bot); pPr.append(pBdr)
p.paragraph_format.space_after = Pt(6)

# Abstract
add_mixed([
    ("Abstract. ", True, False, False),
    ("This report documents a controlled network scanning and vulnerability assessment "
     "exercise conducted entirely within an isolated virtual lab.  A Metasploitable 2 "
     "target was scanned with Nmap using two complementary scan types (SYN and "
     "version-detection), three high-risk attack scenarios were classified against the "
     "MITRE ATT&CK Enterprise framework, and the top vulnerabilities were identified "
     "with OpenVAS, scored with CVSS v3.1, and accompanied by concrete remediation "
     "proposals.  Reflections connect the practical findings to the Week 2 "
     "access-control models and consider how defenders can detect the scans performed.",
     False, False, False),
], space_after=10)



add_heading("1  Introduction and Environment Description")

add_heading("1.1  Objectives", level=2)
add_para(
    "The goals of this module are to: (i) apply host discovery and port-scanning "
    "techniques in a safe lab setting; (ii) interpret scan output in terms of attack "
    "surface; (iii) classify realistic attack scenarios using an established taxonomy; "
    "and (iv) perform a basic vulnerability assessment, prioritise findings by CVSS "
    "score, and propose remediations.",
    space_after=6)

add_heading("1.2  Lab Topology", level=2)
add_para(
    "The lab runs on a single Windows host with Hyper-V, hosting two VMs on an "
    "isolated host-only network (192.168.56.0/24) with no routing to the internet "
    "or any production network.", space_after=4)

add_table(
    headers=["Role", "OS", "IP", "Key services", "Notes"],
    rows=[
        ["Attacker", "Kali Linux 2024.3", "192.168.56.100", "Nmap 7.95, OpenVAS 22", "fully updated"],
        ["Target",   "Metasploitable 2",  "192.168.56.101", "many intentional vulns", "frozen at release"],
    ],
    col_widths=[2.8, 3.5, 3.2, 3.8, 3.0],
)

add_para(
    "The two machines communicate only over the host-only adapter; no traffic crosses "
    "to external networks during any exercise.", space_after=6)

add_heading("1.3  Tool Selection", level=2)
add_mixed([
    ("Nmap", True, False, False),
    (" was chosen for scanning because its open-source code and exhaustive documentation "
     "make output easy to interpret and reproduce; it also supports both the half-open "
     "SYN mode (useful for stealth comparison) and the ", False, False, False),
    ("-sV", False, False, True),
    (" version-detection mode needed for banner grabbing.  ", False, False, False),
    ("OpenVAS", True, False, False),
    (" (Greenbone Community Edition) was selected for vulnerability assessment because "
     "it is free, actively maintained, and maps findings directly to CVE identifiers "
     "and CVSS scores.  All tool interactions were also wrapped in custom Python scripts "
     "(scanner.py, vuln_assessment.py, attack_classification.py) so that parsing and "
     "analysis logic is transparent and reproducible — following the same philosophy "
     "as the Week 2 model simulations.", False, False, False),
], space_after=10)



add_heading("2  Scanning Results and Analysis")

add_heading("2.1  Host Discovery", level=2)
add_mixed([
    ("A ping sweep (", False, False, False),
    ("nmap -sn 192.168.56.0/24", False, False, True),
    (") confirmed two live hosts: ", False, False, False),
    ("192.168.56.100", False, False, True),
    (" (attacker) and ", False, False, False),
    ("192.168.56.101", False, False, True),
    (" (target).  ICMP echo replies were not filtered, which is expected on a host-only "
     "adapter with no firewall policy applied.", False, False, False),
], space_after=6)

add_heading("2.2  Port Scan Results", level=2)

add_mixed([
    ("Scan 1 – SYN scan", True, False, False),
    (" (", False, False, False),
    ("nmap -sS -p- --open -T4", False, False, True),
    ("): sends a TCP SYN, waits for SYN-ACK (open) or RST (closed), then sends RST "
     "to abort the handshake without completing it.  This is faster and generates less "
     "log noise than a full connect.", False, False, False),
], space_after=3)

add_mixed([
    ("Scan 2 – Version scan", True, False, False),
    (" (", False, False, False),
    ("nmap -sV", False, False, True),
    (") on the 12 most interesting ports: completes the TCP handshake and reads service "
     "banners, trading stealth for detail.", False, False, False),
], space_after=6)

add_table(
    headers=["Port", "State", "Service", "Banner / version (from -sV)"],
    rows=[
        ["21/tcp",   "open", "ftp",          "vsftpd 2.3.4"],
        ["22/tcp",   "open", "ssh",          "OpenSSH 4.7p1 Debian 8ubuntu1"],
        ["23/tcp",   "open", "telnet",       "Linux telnetd"],
        ["25/tcp",   "open", "smtp",         "Postfix smtpd"],
        ["80/tcp",   "open", "http",         "Apache httpd 2.2.8 (Ubuntu DAV/2)"],
        ["111/tcp",  "open", "rpcbind",      "2 (RPC #100000)"],
        ["139/tcp",  "open", "netbios-ssn",  "Samba smbd 3.0.20-Debian"],
        ["445/tcp",  "open", "microsoft-ds", "Samba smbd 3.0.20-Debian"],
        ["3306/tcp", "open", "mysql",        "MySQL 5.0.51a-3ubuntu5"],
        ["5432/tcp", "open", "postgresql",   "PostgreSQL 8.3.0–8.3.7"],
        ["6667/tcp", "open", "irc",          "UnrealIRCd 3.2.8.1"],
        ["8180/tcp", "open", "http",         "Apache Tomcat/Coyote JSP 1.1"],
    ],
    col_widths=[2.0, 1.8, 3.2, 9.3],
    mono_cols=[0],
)

add_mixed([
    ("Analysis.  ", True, False, False),
    ("The target exposes an unusually large attack surface: 12 ports open on a machine "
     "that should in practice serve only one or two functions.  Several findings stand out:",
     False, False, False),
], space_after=4)

bullet([("Legacy protocols", True, False, False),
        (": Telnet (port 23) and FTP (port 21) transmit credentials in cleartext — "
         "a passive network tap captures them without any active exploit.", False, False, False)])
bullet([("Outdated versions", True, False, False),
        (": vsftpd 2.3.4 is known to carry a backdoor (CVE-2011-2523); "
         "Samba 3.0.20 has a code-injection flaw (CVE-2007-2447); "
         "Apache 2.2.8 and PHP 5.2 are long past end-of-life.", False, False, False)])
bullet([("Filtered vs. closed", True, False, False),
        (": No ports were returned as filtered in this environment, meaning no firewall "
         "drops packets silently.  In a production setting, filtered ports indicate a "
         "firewall rule rather than an absent service — useful for mapping the defensive "
         "perimeter.", False, False, False)])
bullet([("Detection risk", True, False, False),
        (": The SYN scan is less likely to appear in application logs (the handshake is "
         "never completed), but it is still visible to network IDS tools monitoring for "
         "RST floods or SYN-without-ACK.  The version scan always completes the handshake "
         "and leaves entries in every service's connection log.", False, False, False)])

doc.add_paragraph()


add_heading("3  Attack Classification and Vulnerability Findings")

add_heading("3.1  MITRE ATT&CK Mapping", level=2)
add_para(
    "Three attack scenarios were selected to represent the most realistic threats to the "
    "discovered services.  Each is classified by ATT&CK tactic (the why) and technique "
    "(the how), and additionally mapped to STRIDE to show the security property violated.",
    space_after=4)

add_table(
    headers=["#", "Service", "Tactic (ID)", "Technique", "STRIDE"],
    rows=[
        ["1", "vsftpd 2.3.4 (21)",       "Initial Access / TA0001",   "T1190",     "Elevation of Privilege, Information Disclosure"],
        ["2", "Samba 3.0.20 (445)",       "Execution / TA0002",        "T1059.004", "Tampering, Elevation of Privilege"],
        ["3", "OpenSSH / Debian key (22)","Credential Access / TA0006", "T1110.002", "Spoofing, Elevation of Privilege"],
    ],
    col_widths=[0.6, 3.5, 3.8, 2.5, 5.9],
)

add_mixed([
    ("Scenario 1 — vsftpd Backdoor (TA0001 / T1190).  ", True, False, False),
    ("The compromised 2.3.4 release opens a root shell on port 6200 for any connection "
     "whose username contains ':)'.  This is classified under ", False, False, False),
    ("Initial Access", False, True, False),
    (" because the attacker's primary objective at this stage is gaining a first "
     "foothold.  T1190 (Exploit Public-Facing Application) accurately captures the "
     "mechanism: a network-exposed service is targeted using a publicly known vulnerability.",
     False, False, False),
], space_after=4)

add_mixed([
    ("Scenario 2 — Samba Command Injection (TA0002 / T1059.004).  ", True, False, False),
    ("The ", False, False, False),
    ("username map script", False, False, True),
    (" directive is evaluated by ", False, False, False),
    ("/bin/sh", False, False, True),
    (" without sanitisation.  An attacker injects shell metacharacters to obtain a "
     "reverse shell.  The primary ATT&CK tactic is ", False, False, False),
    ("Execution", False, True, False),
    (" because the adversary's immediate gain is code running on the victim; "
     "T1059.004 (Unix Shell) names the specific interpreter.", False, False, False),
], space_after=4)

add_mixed([
    ("Scenario 3 — Predictable Debian SSH Keys (TA0006 / T1110.002).  ", True, False, False),
    ("CVE-2008-0166 reduced the PRNG seed space on Debian to 32,767 values.  "
     "Iterating through a pre-computed key dictionary to authenticate is classified "
     "under ", False, False, False),
    ("Credential Access", False, True, False),
    (" / Brute Force (T1110), sub-technique T1110.002, because the attacker is "
     "effectively exhausting a small credential space rather than guessing single "
     "passwords.", False, False, False),
], space_after=6)

add_heading("3.2  Top Vulnerability Findings", level=2)

add_table(
    headers=["CVE", "CVSS v3.1", "Severity", "Title"],
    rows=[
        ["CVE-2011-2523", "9.8", "CRITICAL", "vsftpd 2.3.4 Backdoor RCE"],
        ["CVE-2007-2447", "9.8", "CRITICAL", "Samba Username Map Script RCE"],
        ["CVE-2012-1823", "9.8", "CRITICAL", "PHP CGI Argument Injection RCE"],
        ["CVE-2008-0166", "7.8", "HIGH",     "Debian Predictable OpenSSL PRNG"],
        ["CVE-2004-2761", "5.3", "MEDIUM",   "MD5 Collision in SSL Certificates"],
    ],
    col_widths=[3.5, 2.2, 2.5, 8.1],
    mono_cols=[0],
)

add_mixed([
    ("The three CRITICAL findings share the same CVSS base vector ", False, False, False),
    ("AV:N/AC:L/PR:N/UI:N/S:U/C:H/I:H/A:H", False, False, True),
    (": network-reachable, low complexity, no privileges or user interaction required, "
     "and full compromise of confidentiality, integrity, and availability.  This means "
     "any host on the same network (or the internet, if exposed) can obtain root access "
     "with no preconditions — the highest-possible risk.  CVE-2008-0166 scores 7.8 "
     "because it requires local access or prior key enumeration.", False, False, False),
], space_after=10)


add_heading("4  Remediation Proposals and Personal Reflection")

add_heading("4.1  Remediation", level=2)

add_mixed([
    ("CVE-2011-2523 — vsftpd 2.3.4 Backdoor.  ", True, False, False),
    ("Replace vsftpd immediately with version 3.0.5 or later via the distribution "
     "package manager.  If FTP is not operationally required, remove the service and "
     "switch all file-transfer workflows to SFTP (which is already provided by the "
     "OpenSSH daemon on the same host).  Block port 21/tcp at the firewall as a "
     "defence-in-depth measure.", False, False, False),
], space_after=4)

add_mixed([
    ("CVE-2007-2447 — Samba Username Map Script RCE.  ", True, False, False),
    ("Upgrade Samba to 3.0.25 or any current release.  Remove or blank the ", False, False, False),
    ("username map script", False, False, True),
    (" smb.conf directive if it is not actively in use.  Restrict ports 139 and 445 "
     "to trusted subnets via firewall rules; SMB should never be exposed to the "
     "internet.  Consider replacing Samba with a dedicated file-share solution if "
     "the full SMB feature set is not needed.", False, False, False),
], space_after=4)

add_mixed([
    ("CVE-2012-1823 — PHP CGI Argument Injection.  ", True, False, False),
    ("Upgrade PHP to at minimum 5.3.13 / 5.4.3 (or the currently supported 8.x branch).  "
     "Switch the web server from CGI mode to PHP-FPM, which does not pass query strings "
     "to the interpreter binary.  As an immediate workaround, add a rewrite rule to "
     "strip query strings beginning with ", False, False, False),
    ("-", False, False, True),
    (" from PHP script URIs.", False, False, False),
], space_after=6)

add_para("Summary table.", bold=True, space_after=3)
add_table(
    headers=["CVE", "Remediation action"],
    rows=[
        ["CVE-2011-2523", "Upgrade vsftpd; remove FTP; block port 21 at firewall"],
        ["CVE-2007-2447", "Upgrade Samba; clear username map script; restrict ports 139/445"],
        ["CVE-2012-1823", "Upgrade PHP; switch to PHP-FPM; add rewrite rule"],
        ["CVE-2008-0166", "Regenerate all SSH keys; audit authorized_keys; upgrade OpenSSL"],
        ["CVE-2004-2761", "Reissue TLS certificates with SHA-256; disable MD5 cipher suites"],
    ],
    col_widths=[3.5, 12.8],
    mono_cols=[0],
)

add_heading("4.2  Personal Reflection", level=2)

add_para(
    "The most surprising observation was how many critical vulnerabilities are reachable "
    "with zero preconditions.  The vsftpd backdoor requires only a TCP connection and a "
    "username containing two punctuation characters — an attacker with Metasploit can "
    "exploit it in under thirty seconds.  This made the CVSS scoring feel very real "
    "rather than academic: a 9.8 score genuinely means \u201ccompromise is trivially easy "
    "from anywhere on the network.\u201d",
    space_after=6)

add_para(
    "Comparing the two scan types reinforced a principle from the Week 2 access-control "
    "discussion: information has value, and generating it has cost.  The SYN scan is "
    "faster and quieter but tells us only which ports are open.  The version scan costs "
    "a completed TCP handshake per port but reveals the exact software version that maps "
    "to CVE records — the extra information is worth the exposure risk in a controlled "
    "assessment.",
    space_after=6)

add_para(
    "A limitation I encountered is that OpenVAS produces false positives, particularly "
    "for outdated banner-version checks: the scanner reads the version string in a "
    "service banner and matches it against CVE records without verifying whether the "
    "patch has been back-ported.  I cross-referenced each finding against the NVD entry "
    "(nvd.nist.gov) and the Metasploitable package list to confirm the vulnerabilities "
    "were genuine on this specific image.",
    space_after=6)

add_para(
    "Regarding defender detectability: a network IDS such as Snort or Suricata would "
    "detect the SYN scan via a rule matching large numbers of half-open connections to "
    "sequential ports (RST after SYN-ACK pattern).  The version scan is even more "
    "visible because each connection is complete and the banner-grabbing payload is "
    "distinctive.  From a defender's perspective, both scans should trigger an alert "
    "within seconds of starting.",
    space_after=6)

add_para(
    "Connecting to the broader course arc: the formal models of Week 2 describe what "
    "should be enforced; this week reveals what actually runs on a system and how easily "
    "misconfigurations and unpatched software undermine those formal guarantees.  A "
    "Bell-LaPadula policy is meaningless if a root shell is available via a two-character "
    "backdoor in an FTP daemon.  The coming modules on operating-system security and "
    "network defences will, I expect, examine how to layer technical controls so that "
    "even a misconfigured service does not result in total compromise.",
    space_after=10)


# ── REFERENCES ───────────────────────────────────────────────────────────────

add_heading("References")

refs = [
    ([("MITRE Corporation. (2024). ", False, False, False),
      ("ATT&CK Enterprise Matrix v15", False, True, False),
      (". https://attack.mitre.org", False, False, True)]),
    ([("MITRE Corporation. (2024). ", False, False, False),
      ("CVE Programme", False, True, False),
      (". https://www.cve.org", False, False, True)]),
    ([("FIRST. (2023). ", False, False, False),
      ("Common Vulnerability Scoring System v3.1 Specification", False, True, False),
      (". https://www.first.org/cvss", False, False, True)]),
    ([("National Institute of Standards and Technology. (2024). ", False, False, False),
      ("National Vulnerability Database", False, True, False),
      (". https://nvd.nist.gov", False, False, True)]),
    ([("Greenbone Networks. (2024). ", False, False, False),
      ("Greenbone Community Edition / OpenVAS", False, True, False),
      (". https://www.greenbone.net", False, False, True)]),
    ([("Lyon, G. (2024). ", False, False, False),
      ("Nmap Network Scanning: The Official Nmap Project Guide", False, True, False),
      (". https://nmap.org/book/", False, False, True)]),
    ([("Rapid7. (2012). ", False, False, False),
      ("Metasploitable 2 Exploitability Guide", False, True, False),
      (". https://docs.rapid7.com/metasploit/metasploitable-2-exploitability-guide/", False, False, True)]),
    ([("SANS Institute. (2024). ", False, False, False),
      ("TCP/IP and tcpdump", False, True, False),
      (". https://www.sans.org/posters/", False, False, True)]),
    ([("Shostack, A. (2002). ", False, False, False),
      ("STRIDE Threat Model", False, True, False),
      (". Microsoft Security Engineering.", False, False, False)]),
    ([("OWASP Foundation. (2021). ", False, False, False),
      ("OWASP Top Ten 2021", False, True, False),
      (". https://owasp.org/Top10/", False, False, True)]),
]

for i, parts in enumerate(refs, 1):
    ref_item(i, parts)


out = "/home/maria/Computer_Security/Week3_Network_Scanning/report_week3.docx"
doc.save(out)
print(f"saved → {out}")
