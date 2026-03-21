# Builds report_week6.docx from scratch using python-docx
# Run once: python3 build_docx.py
# Output:   report_week6.docx

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)


def set_font(run, name="Calibri", size=11, bold=False, italic=False, color=None):
    run.font.name   = name
    run.font.size   = Pt(size)
    run.font.bold   = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_para(text="", size=11, bold=False, italic=False, align=None,
             space_before=0, space_after=6, keep_with_next=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before   = Pt(space_before)
    p.paragraph_format.space_after    = Pt(space_after)
    p.paragraph_format.line_spacing   = Pt(size * 1.15)
    p.paragraph_format.keep_with_next = keep_with_next
    if align:
        p.alignment = align
    if text:
        run = p.add_run(text)
        set_font(run, size=size, bold=bold, italic=italic)
    return p


def add_heading(text, level=1):
    sizes   = {1: 13, 2: 11}
    space_b = {1: 10, 2: 6}
    p = add_para(text, size=sizes[level], bold=True,
                 space_before=space_b[level], space_after=3,
                 keep_with_next=True)
    if level == 1:
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"),   "single")
        bottom.set(qn("w:sz"),    "4")
        bottom.set(qn("w:space"), "2")
        bottom.set(qn("w:color"), "000000")
        pBdr.append(bottom)
        pPr.append(pBdr)
    return p


def add_mixed(parts, space_before=0, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    p.paragraph_format.line_spacing = Pt(11 * 1.15)
    for text, bold, italic, mono in parts:
        run = p.add_run(text)
        run.font.name   = "Courier New" if mono else "Calibri"
        run.font.size   = Pt(10 if mono else 11)
        run.font.bold   = bold
        run.font.italic = italic
    return p


def add_table(headers, rows, col_widths=None, mono_cols=None):
    mono_cols = mono_cols or []
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style     = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.paragraphs[0].clear()
        run = cell.paragraphs[0].add_run(h)
        set_font(run, size=10, bold=True)

    for ri, row in enumerate(rows):
        trow = table.rows[ri + 1]
        for ci, val in enumerate(row):
            cell = trow.cells[ci]
            cell.paragraphs[0].clear()
            run = cell.paragraphs[0].add_run(str(val))
            set_font(run, name="Courier New" if ci in mono_cols else "Calibri",
                     size=9 if ci in mono_cols else 10)

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    doc.add_paragraph()
    return table


def bullet(parts):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after  = Pt(3)
    p.paragraph_format.line_spacing = Pt(11 * 1.15)
    p.paragraph_format.left_indent  = Cm(0.5)
    for text, bold, italic, mono in parts:
        run = p.add_run(text)
        run.font.name   = "Courier New" if mono else "Calibri"
        run.font.size   = Pt(10 if mono else 11)
        run.font.bold   = bold
        run.font.italic = italic
    return p


def ref_item(num, parts):
    p = doc.add_paragraph()
    p.paragraph_format.space_after       = Pt(2)
    p.paragraph_format.left_indent       = Cm(0.8)
    p.paragraph_format.first_line_indent = Cm(-0.8)
    p.paragraph_format.line_spacing      = Pt(11 * 1.15)
    run = p.add_run(f"[{num}] ")
    set_font(run, size=10, bold=True)
    for text, bold, italic, url in parts:
        run = p.add_run(text)
        run.font.size   = Pt(10)
        run.font.bold   = bold
        run.font.italic = italic
        if url:
            run.font.color.rgb = RGBColor(0x00, 0x56, 0xD2)
    return p


# TITLE

add_para("Report Module 6", size=16, bold=True,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
add_para("Industrial Espionage in Cyberspace", size=13,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
add_para("CS 475: Introduction to Computer Security", size=11,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
add_para("Mária Nyolcas  |  Prof: Asen Nikolov Grozdanov", size=11,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
add_para("March 2026", size=11,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)

p = doc.add_paragraph()
pPr = p._p.get_or_add_pPr()
pBdr = OxmlElement("w:pBdr")
bot = OxmlElement("w:bottom")
bot.set(qn("w:val"), "single"); bot.set(qn("w:sz"), "6")
bot.set(qn("w:space"), "1");    bot.set(qn("w:color"), "000000")
pBdr.append(bot); pPr.append(pBdr)
p.paragraph_format.space_after = Pt(6)

add_mixed([
    ("Abstract.  ", True, False, False),
    ("This report documents a structured passive OSINT investigation of the fictional "
     "target InnovateTech Solutions, alongside a five-phase analysis of the AeroTech "
     "GmbH composite espionage case study. Findings show that publicly available "
     "information alone is sufficient to build a dangerous threat profile. Defensive "
     "countermeasures are proposed for each attack phase, grounded in NIST CSF, CIS "
     "Controls, and ISO/IEC 27001. The report answers all three discussion questions "
     "from the module brief and concludes with reflections connecting the lab evidence "
     "to earlier access-control and vulnerability work.",
     False, False, False),
], space_after=10)


# SECTION 1

add_heading("1  Introduction")

add_para(
    "Industrial espionage is the covert acquisition of confidential business or technical "
    "information for competitive, financial, or strategic advantage. When conducted through "
    "digital networks it becomes cyber-enabled economic espionage \u2014 one of the most "
    "financially damaging threat categories facing organisations today. ENISA consistently "
    "ranks it among the top threats in its annual landscape reports, and the FBI has "
    "described it as \u201cthe greatest long-term threat to the nation\u2019s economy.\u201d",
    space_after=6)

add_para(
    "Unlike ransomware or DDoS attacks, espionage campaigns are designed to remain "
    "invisible. The adversary\u2019s goal is not disruption but quiet, sustained access "
    "to intellectual property, R\u0026D data, and strategic plans. This changes the "
    "defensive calculus: conventional perimeter security detects active attacks; detecting "
    "a patient, low-noise adversary requires different instrumentation, different "
    "monitoring philosophies, and a much longer detection horizon.",
    space_after=6)

add_para(
    "This report covers a passive OSINT investigation of InnovateTech Solutions "
    "(fictional, advanced materials / aerospace R\u0026D), an analysis of the AeroTech "
    "GmbH five-phase case study, and defensive recommendations grounded in established "
    "frameworks. Connections to prior modules are drawn throughout, particularly to the "
    "vulnerability and access-control work of Weeks 2 and 4.",
    space_after=10)


# SECTION 2

add_heading("2  Methodology")

add_para(
    "The investigation followed the six-step lab structure. All work was passive: no "
    "network contact with any real system. Findings are encoded in "
    "osint_recon.py (structured Python dataclasses) and espionage_lifecycle.py "
    "(five-phase case study with ATT\u0026CK mappings and discussion answers).",
    space_after=6)

add_table(
    headers=["Step", "Tool / Source", "Purpose"],
    rows=[
        ["Passive web recon",   "Search operators, WHOIS",    "Domain, registrar, admin contacts, product versions"],
        ["Legacy content",      "Wayback Machine",            "Historical pages revealing old tech versions"],
        ["Code repositories",   "GitHub search",              "Accidentally committed secrets"],
        ["Patent intelligence", "EPO database",               "R&D direction and researcher identities"],
        ["Employee profiling",  "LinkedIn, ResearchGate",     "Roles, skill sets, project names"],
        ["Technical footprint", "dig (DNS), crt.sh",          "Exposed subdomains and services"],
    ],
    col_widths=[3.5, 3.5, 9.3],
)


# SECTION 3

add_heading("3  Findings")

add_heading("3.1  Passive Web Reconnaissance", level=2)

add_para(
    "Five high-value findings emerged without touching the target\u2019s systems. "
    "The most significant was a config.yml file committed to a public GitHub "
    "repository by a junior engineer in November 2022, containing the internal subnet "
    "10.50.0.0/16 and staging hostname sim-lab01.internal. This gives an adversary "
    "internal network topology without any active scanning.",
    space_after=4)

add_para(
    "A 2021 Wayback Machine snapshot disclosed the VPN product and version "
    "(Cisco AnyConnect 4.9, affected by CVE-2021-1247). Job postings named ANSYS "
    "2023 R2 and GitLab CE by version. EPO patent filings identified three active "
    "R\u0026D areas and named the lead researchers by full name.",
    space_after=6)

add_heading("3.2  Employee and Personnel Profiling", level=2)

add_para(
    "Four persons of interest were identified. Dr. K. Brenner (Head of Materials "
    "Research) named the internal project ComposiX-II on LinkedIn. T. Okafor "
    "(Senior IT Administrator) had a GitHub commit referencing a Vault token path. "
    "Dr. A. Reyes (Lead Aerospace Engineer) is a co-inventor on three recent patents "
    "with a public conference biography naming the research lab location. M. Hofer "
    "(Junior R\u0026D Engineer) is the committer of the 2022 secret leak; their email "
    "address is visible in the public git log, making them the easiest initial-access "
    "target via a plausible phishing pretext.",
    space_after=6)

add_heading("3.3  Technical Footprinting", level=2)

add_para(
    "DNS enumeration revealed the self-hosted GitLab instance is internet-exposed "
    "with no IP restriction. The SPF record uses softfail (~all) rather than reject "
    "(-all), meaning spoofed sender addresses will reach most mail servers without "
    "rejection. No DMARC record was found. Certificate transparency logs (crt.sh) "
    "disclosed four subdomains, including sim-lab01 and dev-api, which should not "
    "appear in public certificate logs.",
    space_after=6)

add_heading("3.4  Attack Surface Summary", level=2)

add_table(
    headers=["Finding", "Adversarial value", "Remediation"],
    rows=[
        ["Leaked subnet + hostname (GitHub)", "Internal recon without network contact",   "Secrets scanning pre-commit hook (gitleaks)"],
        ["No DMARC / SPF softfail",           "Spoofed-sender phishing at scale",         "Add DMARC p=reject; harden SPF to -all"],
        ["GitLab internet-exposed",            "Unauthenticated code access",              "Move behind VPN or IP allowlist"],
        ["VPN version disclosed (Wayback)",    "CVE-targeted exploit",                     "Strip version strings from all public content"],
        ["Researcher named on patents",        "Targeted spear-phishing pretext",          "Omit internal project names from public bios"],
    ],
    col_widths=[4.5, 5.0, 6.8],
)


# SECTION 4

add_heading("4  Analysis")

add_heading("4.1  AeroTech Case Study Discussion Questions", level=2)

add_mixed([
    ("Q1 \u2014 Best disruption point.  ", True, False, False),
    ("Phase 2 (Initial Access via spear-phishing) is the single highest-leverage "
     "intervention point. Before the PDF exploit executes, the attacker has no foothold "
     "and no intelligence about internal systems beyond what passive OSINT provided. "
     "Blocking delivery \u2014 through email sandboxing, mandatory patch cycles for "
     "document viewers, and phishing simulation training \u2014 collapses the entire "
     "chain. At every subsequent phase the adversary already has a foothold, and defence "
     "becomes progressively more expensive and less certain.",
     False, False, False),
], space_after=6)

add_mixed([
    ("Q2 \u2014 Least privilege and Phase 3.  ", True, False, False),
    ("The lateral movement in Phase 3 succeeded because shared service accounts had "
     "broad access across network segments. If the procurement workstation had been "
     "separated from the R\u0026D file server by a firewall requiring explicit "
     "authorisation \u2014 enforcing the least-privilege principle examined in Week 2 "
     "\u2014 the attacker would have stalled at the first segment boundary. The blast "
     "radius of the initial compromise would have been limited to a single low-value "
     "host even after successful initial access.",
     False, False, False),
], space_after=6)

add_mixed([
    ("Q3 \u2014 Human behaviour across all phases.  ", True, False, False),
    ("No phase was purely technical. Phase 1: a junior engineer committed secrets to a "
     "public repository. Phase 2: a researcher opened an attachment without verifying "
     "the sender domain. Phase 3: administrators shared service accounts for convenience. "
     "Phase 5: detection relied on a single analyst reviewing historical logs during a "
     "manual audit. Technology failed at every step because it was misconfigured or "
     "absent \u2014 the root cause in each case was a human decision.",
     False, False, False),
], space_after=8)

add_heading("4.2  DLP Policy Requirements", level=2)

add_para(
    "A DLP policy addressing the identified risks needs three layers. At the endpoint: "
    "block bulk copy of files tagged Confidential or R\u0026D to removable media or "
    "personal cloud storage. At the network boundary: alert on outbound transfers "
    "exceeding a volume threshold to cloud domains not on an approved list; inspect TLS "
    "for anomalous certificate chains. At the repository: pre-receive hooks rejecting "
    "commits matching patterns for IP ranges, hostnames, or secret tokens.",
    space_after=6)

add_heading("4.3  Highest Impact-to-Cost Defensive Control", level=2)

add_para(
    "Across all findings, secrets scanning in the CI/CD pipeline has the highest "
    "impact-to-cost ratio. A pre-commit hook running gitleaks costs near zero to deploy, "
    "requires no ongoing licensing, and would have prevented the single most actionable "
    "finding in this investigation \u2014 the leaked subnet and hostname \u2014 before "
    "it ever reached a public repository. Every other finding derived part of its value "
    "from that one commit.",
    space_after=10)


# CONCLUSION (unnumbered)

add_heading("Conclusion")

add_para(
    "This module demonstrated that the most dangerous phase of an espionage campaign "
    "requires no hacking at all: passive OSINT alone produced five high-value findings "
    "against a fictional target, including internal network topology, CVE-targetable "
    "software versions, and direct candidate identities for spear-phishing. The AeroTech "
    "case study showed that human decisions created every exploitable gap, and that the "
    "highest-leverage defensive investment is therefore in people and process, not "
    "exclusively in technology.",
    space_after=6)

add_para(
    "Looking ahead to network and operating-system security, industrial espionage makes "
    "clear why layered defence matters: an attacker who spends four weeks moving "
    "laterally undetected has effectively unlimited time to find and extract the crown "
    "jewels. The goal of layered defence is to ensure that no single failure \u2014 "
    "human or technical \u2014 is sufficient to reach them.",
    space_after=10)


# REFERENCES

add_heading("References")

refs = [
    ([("Center for Internet Security. (2024). ", False, False, False),
      ("CIS Controls v8", False, True, False),
      (". https://www.cisecurity.org/controls", False, False, True)]),
    ([("Council of Europe. (2001). ", False, False, False),
      ("Budapest Convention on Cybercrime", False, True, False),
      (". https://www.coe.int/en/web/cybercrime/the-budapest-convention", False, False, True)]),
    ([("European Parliament. (2016). ", False, False, False),
      ("Directive 2016/943 on the protection of trade secrets", False, True, False),
      (". https://eur-lex.europa.eu", False, False, True)]),
    ([("European Union Agency for Cybersecurity [ENISA]. (2024). ", False, False, False),
      ("ENISA threat landscape 2024", False, True, False),
      (". https://www.enisa.europa.eu/publications/enisa-threat-landscape", False, False, True)]),
    ([("Federal Bureau of Investigation. (2023). ", False, False, False),
      ("Economic espionage", False, True, False),
      (". https://www.fbi.gov/investigate/counterintelligence/economic-espionage", False, False, True)]),
    ([("IP Commission. (2017). ", False, False, False),
      ("The theft of American intellectual property: Reassessments of the challenge and United States policy", False, True, False),
      (". https://www.ipcommission.org", False, False, True)]),
    ([("Lockheed Martin. (2022). ", False, False, False),
      ("Cyber kill chain", False, True, False),
      (". https://www.lockheedmartin.com/en-us/capabilities/cyber/cyber-kill-chain.html", False, False, True)]),
    ([("MITRE Corporation. (2024). ", False, False, False),
      ("ATT&CK enterprise matrix v15", False, True, False),
      (". https://attack.mitre.org", False, False, True)]),
    ([("National Institute of Standards and Technology [NIST]. (2018). ", False, False, False),
      ("Framework for improving critical infrastructure cybersecurity", False, True, False),
      (" (v1.1). https://www.nist.gov/cyberframework", False, False, True)]),
    ([("SANS Institute. (2024). ", False, False, False),
      ("Reading room: Insider threat and data exfiltration", False, True, False),
      (". https://www.sans.org/reading-room/", False, False, True)]),
]

for i, parts in enumerate(refs, 1):
    ref_item(i, parts)


OUT = "/home/maria/Computer_Security/Week6_Industrial_Espionage/report_week6.docx"
doc.save(OUT)
print(f"saved -> {OUT}")
