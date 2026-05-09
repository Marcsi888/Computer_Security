# Builds report_week11.docx from scratch using python-docx
# Run once: python build_docx.py
# Output:   report_week11.docx

import os
from docx import Document

_HERE = os.path.dirname(os.path.abspath(__file__))
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)


def set_font(run, name="Calibri", size=11, bold=False, italic=False, color=None):
    run.font.name   = name
    run.font.size   = Pt(size)
    run.font.bold   = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_para(text="", size=11, bold=False, italic=False, align=None,
             space_before=0, space_after=6, keep_with_next=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before   = Pt(space_before)
    p.paragraph_format.space_after    = Pt(space_after)
    p.paragraph_format.line_spacing   = Pt(size * 1.15)
    p.paragraph_format.keep_with_next = keep_with_next
    if align:
        p.alignment = align
    if text:
        run = p.add_run(text)
        set_font(run, size=size, bold=bold, italic=italic)
    return p


def add_heading(text, level=1):
    sizes   = {1: 13, 2: 11}
    space_b = {1: 10, 2: 6}
    p = add_para(text, size=sizes[level], bold=True,
                 space_before=space_b[level], space_after=3,
                 keep_with_next=True)
    if level == 1:
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"),   "single")
        bottom.set(qn("w:sz"),    "4")
        bottom.set(qn("w:space"), "2")
        bottom.set(qn("w:color"), "000000")
        pBdr.append(bottom)
        pPr.append(pBdr)
    return p


def add_mixed(parts, space_before=0, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    p.paragraph_format.line_spacing = Pt(11 * 1.15)
    for text, bold, italic, mono in parts:
        run = p.add_run(text)
        run.font.name   = "Courier New" if mono else "Calibri"
        run.font.size   = Pt(10 if mono else 11)
        run.font.bold   = bold
        run.font.italic = italic
    return p


def add_table(headers, rows, col_widths=None, mono_cols=None):
    mono_cols = mono_cols or []
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style     = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.paragraphs[0].clear()
        run = cell.paragraphs[0].add_run(h)
        set_font(run, size=10, bold=True)

    for ri, row in enumerate(rows):
        trow = table.rows[ri + 1]
        for ci, val in enumerate(row):
            cell = trow.cells[ci]
            cell.paragraphs[0].clear()
            run = cell.paragraphs[0].add_run(str(val))
            set_font(run, name="Courier New" if ci in mono_cols else "Calibri",
                     size=9 if ci in mono_cols else 10)

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    doc.add_paragraph()
    return table


def bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after  = Pt(3)
    p.paragraph_format.line_spacing = Pt(11 * 1.15)
    p.paragraph_format.left_indent  = Cm(0.5)
    if bold_prefix:
        run = p.add_run(bold_prefix)
        set_font(run, bold=True)
    run = p.add_run(text)
    set_font(run)
    return p


# ---------------------------------------------------------------------------
# TITLE BLOCK
# ---------------------------------------------------------------------------

add_para("Report Module 11", size=16, bold=True,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
add_para("Computer Security Software", size=13,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
add_para("CS 475: Introduction to Computer Security", size=11,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
add_para("Maria Nyolcas  |  Prof: Asen Nikolov Grozdanov", size=11,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
add_para("May 2026", size=11,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)

p = doc.add_paragraph()
pPr = p._p.get_or_add_pPr()
pBdr = OxmlElement("w:pBdr")
bot = OxmlElement("w:bottom")
bot.set(qn("w:val"), "single"); bot.set(qn("w:sz"), "6")
bot.set(qn("w:space"), "1");    bot.set(qn("w:color"), "000000")
pBdr.append(bot); pPr.append(pBdr)
p.paragraph_format.space_after = Pt(6)

add_mixed([
    ("Abstract.  ", True, False, False),
    ("This report covers the Week 11 deliverable of the continuous security portfolio: the "
     "landscape of computer security software, a hands-on evaluation of Windows Defender "
     "against NIST SP\u00a0800-83 and CIS Control\u00a010 criteria, and a prioritised advisory "
     "recommendation for a small university department. Three EICAR-based detection tests "
     "and a Windows Firewall rule audit form the lab evidence base. All six advisory "
     "recommendations are implementable with free or open-source tools. The synthesis "
     "connects this week\u2019s findings to the operational gaps identified across the "
     "earlier portfolio modules.",
     False, False, False),
], space_after=10)


# ---------------------------------------------------------------------------
# SECTION 1 - INTRODUCTION
# ---------------------------------------------------------------------------

add_heading("1  Introduction")

add_para(
    "A common mistake among newcomers to cybersecurity is imagining that installing the right "
    "software solves the problem. It does not. Security software is one layer of defence in "
    "depth \u2014 the principle that no single control, however capable, should be the last "
    "line of defence. The NIST Cybersecurity Framework organises controls into Identify, "
    "Protect, Detect, Respond, and Recover, and software tools appear across every one of "
    "those functions, working alongside policy, people, and physical controls.",
    space_after=6)

add_para(
    "This week\u2019s activities move from understanding what security mechanisms exist to "
    "evaluating how software implements them, why some implementations fail, and what a "
    "small organisation should prioritise. The report is structured into four parts: a "
    "theoretical taxonomy of the eight major security software families, a hands-on tool "
    "evaluation and lab, a non-technical advisory report, and a reflective synthesis "
    "connecting findings to earlier portfolio modules.",
    space_after=6)


# ---------------------------------------------------------------------------
# SECTION 2 - SECURITY SOFTWARE LANDSCAPE (THEORY)
# ---------------------------------------------------------------------------

add_heading("2  Security Software Landscape")

add_heading("2.1  Major Categories", level=2)

add_para(
    "Eight major families of security software are summarised below, each mapped to the "
    "primary threat it addresses, its deployment layer, and its alignment with the NIST CSF "
    "functions. No single category covers all five CSF functions, which underscores the "
    "defence-in-depth requirement.",
    space_after=4)

add_table(
    headers=["Category", "Acronym", "NIST CSF", "Key limitation"],
    rows=[
        ("AV / Endpoint Detection & Response", "AV/EDR",
         "Protect, Detect, Respond",
         "Signature lag; behavioural engines require cloud connectivity"),
        ("Firewall (host-based and network)", "FW",
         "Protect",
         "Misconfigured rules create a false sense of protection"),
        ("Intrusion Detection / Prevention", "IDS/IPS",
         "Detect, Respond",
         "Rule-tuning burden; encrypted traffic limits deep inspection"),
        ("VPN client / gateway", "VPN",
         "Protect",
         "Protects the tunnel only; misconfigured split-tunnelling leaks traffic"),
        ("Security Information & Event Mgmt", "SIEM",
         "Identify, Detect, Respond, Recover",
         "Dependent on log quality; high alert volume causes analyst fatigue"),
        ("Data Loss Prevention", "DLP",
         "Protect, Detect",
         "Relies on accurate data classification; mislabelled data evades controls"),
        ("Vulnerability Scanner", "VA",
         "Identify",
         "Point-in-time snapshot; misses zero-days and unscanned assets"),
        ("Privileged Access Management", "PAM",
         "Protect, Detect",
         "Only covers enrolled accounts; shadow admin accounts bypass controls"),
    ],
    col_widths=[5.2, 1.6, 3.8, 5.6],
)

add_heading("2.2  Defence-in-Depth Layer Mapping", level=2)

add_para(
    "The five-layer model below shows where each software family sits in a defence-in-depth "
    "architecture. The key insight is that each layer assumes the previous one has been "
    "bypassed and provides independent detection or containment capability.",
    space_after=4)

add_table(
    headers=["Layer", "Software families", "Rationale"],
    rows=[
        ("Perimeter",         "NGFW, inline IPS, VPN gateway",
         "Controls all ingress and egress; first barrier attackers must cross."),
        ("Network internal",  "IDS (passive tap), SIEM ingestion, DLP proxy",
         "Detects lateral movement and exfiltration after a perimeter bypass."),
        ("Host",              "AV/EDR, host FW, HIDS, VA scanner (authenticated)",
         "Last-resort detection at the point where malicious code executes."),
        ("Identity",          "PAM, MFA via IAM tooling",
         "Limits credential-based lateral movement and insider escalation."),
        ("Data",              "DLP agent, encryption at rest",
         "Protects the asset even when all outer layers are compromised."),
    ],
    col_widths=[2.8, 5.4, 8.0],
)


# ---------------------------------------------------------------------------
# SECTION 3 - TOOL EVALUATION AND LAB
# ---------------------------------------------------------------------------

add_heading("3  Tool Evaluation and Hands-on Lab")

add_heading("3.1  Windows Defender Evaluation Matrix", level=2)

add_para(
    "Windows Defender (Microsoft Defender Antivirus) was evaluated against eight criteria "
    "drawn from NIST SP\u00a0800-83 "
    "(\u2018Guide to Malware Incident Prevention and Handling\u2019) and CIS Control\u00a010 "
    "(\u2018Malware Defences\u2019). This tool was selected because it is the default endpoint "
    "protection layer on Windows, requiring no additional installation, and represents the "
    "sole AV control in many small organisations.",
    space_after=4)

add_table(
    headers=["Criterion", "Rating", "Evidence and recommendation"],
    rows=[
        ("Signature coverage", "Meets",
         "Definitions updated multiple times daily; EICAR detected immediately on file-write. "
         "Validate update policy is not delayed by GPO."),
        ("Behavioural / heuristic detection", "Partial",
         "Cloud heuristics (MAPS) available but require internet connectivity; air-gapped "
         "hosts lose this layer. Enable MAPS and automatic sample submission."),
        ("Real-time (on-access) protection", "Meets",
         "EICAR blocked on file-write without manual scan trigger; Event ID\u00a01116 logged. "
         "Enforce via GPO: DisableRealtimeMonitoring\u00a0=\u00a00."),
        ("Scheduled full scan", "Partial",
         "Default schedule is configurable by users. Enforce centrally via Intune or GPO; "
         "alert on missed scans."),
        ("Logging / SIEM integration", "Partial",
         "Logs to Windows Event Log (IDs 1116, 1117, 5001); WEF forwarding requires explicit "
         "configuration. Configure WEF for the Defender Operational log."),
        ("Tamper protection", "Meets",
         "Service stop blocked for unprivileged callers; Access Denied returned on "
         "sc stop WinDefend. Verify via Get-MpPreference | Select TamperProtection."),
        ("Ransomware / Controlled Folder Access", "Partial",
         "Feature present but disabled by default. Enable and allow-list trusted apps "
         "before department-wide rollout."),
        ("Incident response integration", "Gap",
         "Standalone Defender lacks automated playbooks or network isolation without paid "
         "Defender for Endpoint. Supplement with Wazuh + scripted isolation response."),
    ],
    col_widths=[4.2, 1.8, 10.2],
)

add_heading("3.2  Hands-on Detection Tests", level=2)

add_mixed([
    ("Test 1 \u2014 EICAR standard test file.  ", True, False, False),
    ("The EICAR test file was written to the Desktop of a Windows\u00a011 host with "
     "real-time protection active. The file was removed within approximately two seconds "
     "of creation. A Windows Security notification appeared reading \u2018Threat found: "
     "Virus:DOS/EICAR_Test_File\u2019 and Event ID\u00a01116 was recorded in the "
     "Microsoft-Windows-Windows Defender/Operational log with action\u00a0\u2018Quarantine\u2019. "
     "Result: ", False, False, False),
    ("PASS.", True, False, False),
], space_after=6)

add_mixed([
    ("Test 2 \u2014 EICAR inside a ZIP archive.  ", True, False, False),
    ("The EICAR string was packaged into a .zip archive before writing to disk. The archive "
     "was not flagged on creation. When extraction was attempted, Defender blocked "
     "decompression and logged Event ID\u00a01116 with the compressed file path. This confirms "
     "Defender scans compressed content at the point of access rather than on archive "
     "creation. Result: ", False, False, False),
    ("PASS.", True, False, False),
], space_after=6)

add_mixed([
    ("Test 3 \u2014 Tamper protection (service stop attempt).  ", True, False, False),
    ("The command ", False, False, False),
    ("sc stop WinDefend", False, False, True),
    (" was executed from a standard user account. The command returned \u2018Access Denied\u2019 "
     "(error 5) and no Event ID\u00a05001 was generated, confirming that Tamper Protection "
     "prevented service termination. Result: ", False, False, False),
    ("PASS.", True, False, False),
], space_after=6)

add_heading("3.3  Windows Firewall Rule Audit", level=2)

add_para(
    "A review of the default Windows Firewall rule set identified four findings relevant "
    "to the lab environment. Two rules represent actionable risk, one is an acceptable "
    "operational default, and one demonstrates a correct secure-by-default posture.",
    space_after=4)

add_table(
    headers=["Direction", "Rule", "State", "Action", "Risk commentary"],
    rows=[
        ("Inbound", "Remote Desktop (TCP 3389)", "Enabled", "Allow",
         "Open on all profiles; brute-force risk if internet-reachable. Scope to management IPs."),
        ("Inbound", "File and Printer Sharing (SMB, TCP 445)", "Enabled", "Allow",
         "Acceptable on domain profile; widens lateral-movement surface on public profile."),
        ("Outbound", "All outbound (default allow)", "Enabled", "Allow",
         "Malware with C2 beaconing is unblocked. An outbound allow-list limits exfiltration."),
        ("Inbound", "ICMP Echo Request (ping)", "Disabled", "Block",
         "Correct default on public profile; reduces host discovery on untrusted networks."),
    ],
    col_widths=[2.0, 4.4, 1.6, 1.6, 6.6],
)


# ---------------------------------------------------------------------------
# SECTION 4 - ADVISORY REPORT
# ---------------------------------------------------------------------------

add_heading("4  Advisory Report")

add_para(
    "For: Head of Department, University Computing Services",
    italic=True, space_after=2)
add_para(
    "Re: Minimum-viable security software stack for a department of approximately 50 users",
    italic=True, space_after=6)

add_para(
    "The six recommendations below are ordered by priority. Each is implementable with "
    "free or open-source tools at low or medium effort. Items 1, 2, and 6 can be "
    "deployed immediately using software already present in a standard Windows environment "
    "and address the three most common initial-access and persistence patterns seen in "
    "public breach reports: unmonitored endpoint alerts, over-permissive firewall "
    "defaults, and shared credentials.",
    space_after=4)

add_table(
    headers=["#", "Control area", "Gap", "Recommended action", "NIST CSF", "Effort"],
    rows=[
        ("1", "Endpoint protection",
         "Defender unmanaged; Controlled Folder Access disabled",
         "Enforce Tamper Protection, CFA, and daily-update policy via GPO. "
         "Forward Event IDs 1116/1117 to central log store.",
         "Protect / Detect", "Low"),
        ("2", "Firewall hardening",
         "RDP inbound open on all profiles; outbound default-allow",
         "Restrict inbound RDP to management VLAN. Implement outbound allow-list "
         "for known applications.",
         "Protect", "Low"),
        ("3", "Centralised logging",
         "No log aggregation; events isolated per host",
         "Deploy Wazuh or Elastic Security. Alert on Event IDs 4625, 4720, 1116.",
         "Detect / Respond", "Medium"),
        ("4", "Vulnerability management",
         "Patching occurs but no scan-based gap verification",
         "Monthly authenticated OpenVAS scans. CVSS 9+ patched within 72 h; "
         "CVSS 7+ within 30 days.",
         "Identify", "Medium"),
        ("5", "Secure remote access",
         "Direct RDP over internet; no VPN policy",
         "Deploy WireGuard VPN; disable external RDP; enforce MFA on VPN auth.",
         "Protect", "Medium"),
        ("6", "Privileged access",
         "Shared local admin credentials; no session audit",
         "Deploy Microsoft LAPS for unique local admin passwords; enforce just-in-time "
         "elevation logged to SIEM.",
         "Protect / Detect", "Low"),
    ],
    col_widths=[0.6, 2.8, 3.4, 5.4, 2.4, 1.6],
)


# ---------------------------------------------------------------------------
# SECTION 5 - SYNTHESIS AND REFLECTION
# ---------------------------------------------------------------------------

add_heading("5  Synthesis and Reflection")

add_para(
    "The week\u2019s activities reinforce a pattern visible across the entire portfolio: the "
    "most impactful security failures are not exotic zero-days but avoidable configuration "
    "defaults left unchanged. Windows Defender passed all three detection tests without "
    "modification, yet its value to the organisation is severely limited without centralised "
    "management, logging, and response integration. A firewall with a default-allow outbound "
    "rule and open inbound RDP is structurally equivalent to no meaningful perimeter control.",
    space_after=6)

add_para(
    "Connecting this week to earlier portfolio threads: Week\u00a06 demonstrated how OSINT "
    "enables targeted phishing as an initial-access vector; Week\u00a07 showed how exposed "
    "services and missing input validation provide a foothold and persistence path. Computer "
    "security software is the operational layer that would intercept those threats at "
    "execution time \u2014 but only when correctly configured, monitored, and integrated. "
    "A correctly updated Defender instance catches common payload delivery. A SIEM ingesting "
    "Defender logs alerts on lateral movement. A VPN gateway with MFA eliminates direct "
    "credential attacks on management interfaces.",
    space_after=6)

add_para(
    "The recurring observation is that the weakest link is not usually the absence of a tool "
    "but the absence of a policy that enforces, monitors, and validates that tool\u2019s "
    "correct operation. Technology without governance is incomplete security.",
    space_after=10)


# ---------------------------------------------------------------------------
# SAVE
# ---------------------------------------------------------------------------

doc.save(os.path.join(_HERE, "report_week11.docx"))
print("report_week11.docx written successfully.")
